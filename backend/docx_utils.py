"""
Utility functions for extracting text from DOCX files
"""
import zipfile
import xml.etree.ElementTree as ET

def extract_text_from_docx(docx_path):
    """
    Extract text from a DOCX file using a Lambda-compatible approach
    
    Args:
        docx_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text from the document
    """
    try:
        # First try the standard python-docx approach
        try:
            from docx import Document
            doc = Document(docx_path)
            
            # Extract text from paragraphs
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content)
            
        except ImportError:
            # Fallback to manual XML parsing if python-docx fails
            return extract_text_from_docx_manual(docx_path)
        
    except Exception as e:
        # If all else fails, try the manual approach
        try:
            return extract_text_from_docx_manual(docx_path)
        except:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")

def extract_text_from_docx_manual(docx_path):
    """
    Manually extract text from DOCX by parsing the XML structure
    """
    text_content = []
    
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        # Read the main document XML
        doc_xml = docx_zip.read('word/document.xml')
        
        # Parse the XML
        root = ET.fromstring(doc_xml)
        
        # Define namespace
        namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # Extract text from paragraphs
        for para in root.findall('.//w:p', namespace):
            para_text = []
            for text_elem in para.findall('.//w:t', namespace):
                if text_elem.text:
                    para_text.append(text_elem.text)
            
            if para_text:
                text_content.append(''.join(para_text))
    
    return '\n\n'.join(text_content)
